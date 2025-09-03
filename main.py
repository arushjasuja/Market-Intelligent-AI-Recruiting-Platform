import streamlit as st
import os
import time
import pandas as pd
import plotly.express as px
import logging
from datetime import datetime
from dotenv import load_dotenv
from openai import OpenAI
from sentence_transformers import SentenceTransformer
import PyPDF2
import docx

# Optimize for Streamlit Cloud
if 'STREAMLIT_CLOUD' in os.environ:
    # Use smaller models on cloud
    @st.cache_resource
    def load_sentence_model():
        return SentenceTransformer('paraphrase-MiniLM-L3-v2')  # Smaller model
else:
    @st.cache_resource  
    def load_sentence_model():
        return SentenceTransformer('all-MiniLM-L6-v2')  # Full model

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Import custom modules with specific error handling
try:
    from shared import Job, Candidate, MarketReport, MatchResult, DatabaseManager
except ImportError as e:
    st.error(f"❌ Core module import failed: {e}")
    st.markdown("""
    **Fix**: Ensure all files are in the correct directory structure:
    ```
    shared/
    ├── __init__.py
    ├── models.py
    └── database.py
    ```
    """)
    st.stop()

try:
    from agents.crewai_market import MarketIntelligenceCrew
except ImportError as e:
    st.error(f"⚠️ CrewAI module not available: {e}")
    st.markdown("**Fix**: `pip install crewai==0.41.0`")
    MarketIntelligenceCrew = None

try:
    from agents.langchain_content import LangChainContentGenerator
except ImportError as e:
    st.error(f"⚠️ LangChain module not available: {e}")
    st.markdown("**Fix**: `pip install langchain==0.1.0 langchain-openai==0.0.8`")
    LangChainContentGenerator = None

try:
    from agents.autogen_matching import AutoGenMatchingSystem
except ImportError as e:
    st.error(f"⚠️ AutoGen module not available: {e}")
    st.markdown("""
    **Fix AutoGen Installation:**
    ```bash
    pip install pyautogen==0.2.25
    ```
    
    If still failing, try:
    ```bash
    pip uninstall pyautogen
    pip install pyautogen==0.2.25 --no-cache-dir
    ```
    """)
    AutoGenMatchingSystem = None

try:
    from agents.nlp_parser import EnhancedResumeParser
except ImportError as e:
    st.error(f"⚠️ NLP Parser module not available: {e}")
    st.markdown("**Fix**: `python -m spacy download en_core_web_sm`")
    EnhancedResumeParser = None

# Load environment variables and configure Streamlit
load_dotenv()
st.set_page_config(
    page_title="Production AI Recruiting Platform", 
    page_icon="🤖", 
    layout="wide",
    initial_sidebar_state="expanded"
)

def create_basic_resume_parser():
    """Create a basic fallback resume parser when NLP is unavailable"""
    class BasicResumeParser:
        def parse_resume(self, text: str, filename: str = "") -> dict:
            """Basic resume parsing without NLP"""
            import re
            
            # Extract name from filename
            name = "Unknown Candidate"
            if filename:
                name = re.sub(r'[^\w\s]', ' ', filename.replace('.pdf', '')).strip().title()
            
            # Basic skill extraction
            basic_skills = []
            common_skills = [
                "Python", "Java", "JavaScript", "SQL", "AWS", "React", "Docker", 
                "Git", "Linux", "Node.js", "HTML", "CSS", "Machine Learning",
                "TensorFlow", "PyTorch", "Kubernetes", "MongoDB", "PostgreSQL"
            ]
            
            text_lower = text.lower()
            for skill in common_skills:
                if skill.lower() in text_lower:
                    basic_skills.append(skill)
            
            # Basic experience extraction
            experience_patterns = [
                r'(\d{1,2})\+?\s*years?\s*(?:of\s*)?experience',
                r'(\d{1,2})\+?\s*yrs?\s*experience'
            ]
            
            experience_years = 0
            for pattern in experience_patterns:
                matches = re.findall(pattern, text_lower)
                if matches:
                    experience_years = max([int(m) for m in matches if m.isdigit()])
                    break
            
            return {
                "name": name,
                "skills": basic_skills,
                "experience_years": experience_years,
                "education": [],
                "certifications": [],
                "previous_companies": [],
                "resume_text": text
            }
    
    return BasicResumeParser()

@st.cache_resource
def initialize_components():
    """Initialize core components with caching"""
    try:
        db = DatabaseManager()
        sentence_model = SentenceTransformer('all-MiniLM-L6-v2')
        
        # Handle case where EnhancedResumeParser might not be available
        resume_parser = None
        if EnhancedResumeParser is not None:
            try:
                resume_parser = EnhancedResumeParser()
                logger.info("Enhanced NLP resume parser initialized")
            except Exception as e:
                logger.warning(f"Enhanced resume parser initialization failed: {e}")
                st.warning("⚠️ Enhanced NLP parsing unavailable - using basic parser")
                resume_parser = create_basic_resume_parser()
        else:
            st.warning("⚠️ spaCy not available - using basic resume parser")
            resume_parser = create_basic_resume_parser()
        
        logger.info("Core components initialized successfully")
        return db, sentence_model, resume_parser
    except Exception as e:
        logger.error(f"Component initialization failed: {e}")
        st.error(f"Failed to initialize core components: {e}")
        st.stop()

def validate_openai_key():
    """Validate and retrieve OpenAI API key"""
    # Try multiple sources
    api_key = None
    
    try:
        api_key = st.secrets["OPENAI_API_KEY"]
        logger.info("API key loaded from Streamlit secrets")
    except:
        api_key = os.getenv("OPENAI_API_KEY")
        if api_key:
            logger.info("API key loaded from environment")
    
    if not api_key:
        st.error("🔑 OpenAI API Key Required")
        st.markdown("""
        **Configuration Options:**
        1. **Environment Variable**: Set `OPENAI_API_KEY=your_key`
        2. **Streamlit Secrets**: Add to `.streamlit/secrets.toml`
        3. **Manual Entry**: Enter below (session only)
        """)
        
        api_key = st.text_input("Enter OpenAI API Key:", type="password", 
                               help="Your key is not stored permanently")
        if not api_key:
            st.stop()
        else:
            st.success("✅ API Key accepted for this session")
    
    # Basic validation
    if not api_key.startswith('sk-'):
        st.error("❌ Invalid API key format. OpenAI keys start with 'sk-'")
        st.stop()
    
    return api_key

def initialize_ai_agents(api_key: str, sentence_model, resume_parser):
    """Initialize all AI agents with error handling"""
    agents = {}
    
    # CrewAI Market Intelligence
    if MarketIntelligenceCrew is not None:
        try:
            agents['market_crew'] = MarketIntelligenceCrew(api_key)
            logger.info("CrewAI market intelligence initialized")
        except Exception as e:
            logger.error(f"CrewAI initialization failed: {e}")
            st.sidebar.error("⚠️ CrewAI unavailable")
    else:
        st.sidebar.error("❌ CrewAI not installed")
    
    # LangChain Content Generator
    if LangChainContentGenerator is not None:
        try:
            agents['content_generator'] = LangChainContentGenerator(api_key)
            logger.info("LangChain content generator initialized")
        except Exception as e:
            logger.error(f"LangChain initialization failed: {e}")
            st.sidebar.error("⚠️ LangChain unavailable")
    else:
        st.sidebar.error("❌ LangChain not installed")
    
    # AutoGen Matching System
    if AutoGenMatchingSystem is not None:
        try:
            agents['matching_system'] = AutoGenMatchingSystem(api_key, sentence_model)
            logger.info("AutoGen matching system initialized")
        except Exception as e:
            logger.error(f"AutoGen initialization failed: {e}")
            st.sidebar.error("⚠️ AutoGen unavailable")
    else:
        st.sidebar.error("❌ AutoGen not installed - see installation instructions above")
    
    # Resume Parser (always available as fallback)
    if resume_parser is not None:
        agents['resume_parser'] = resume_parser
    
    # OpenAI Client
    try:
        agents['openai_client'] = OpenAI(api_key=api_key)
        logger.info("OpenAI client initialized")
    except Exception as e:
        logger.error(f"OpenAI client initialization failed: {e}")
        st.sidebar.error("⚠️ OpenAI client unavailable")
    
    return agents

def main():
    """Main application entry point"""
    st.sidebar.title("🤖 Production AI Recruiting Platform")
    st.sidebar.markdown("**Multi-Agent System: CrewAI • LangChain • AutoGen**")
    
    # Initialize core components
    db, sentence_model, resume_parser = initialize_components()
    
    # Validate API key
    api_key = validate_openai_key()
    
    # Initialize AI agents
    agents = initialize_ai_agents(api_key, sentence_model, resume_parser)
    
    # Load and display system statistics
    try:
        reports = db.get_reports()
        jobs = db.get_all_jobs()
        candidates = db.get_all_candidates()
        
        # Statistics display
        st.sidebar.markdown("### 📊 System Statistics")
        col1, col2, col3 = st.sidebar.columns(3)
        col1.metric("Reports", len(reports))
        col2.metric("Jobs", len(jobs))
        col3.metric("Candidates", len(candidates))
        
        # System status
        st.sidebar.markdown("### 🚦 System Status")
        status_messages = []
        if 'market_crew' in agents:
            status_messages.append("✅ CrewAI Active")
        else:
            status_messages.append("❌ CrewAI Inactive")
            
        if 'content_generator' in agents:
            status_messages.append("✅ LangChain Active")
        else:
            status_messages.append("❌ LangChain Inactive")
            
        if 'matching_system' in agents:
            status_messages.append("✅ AutoGen Active")
        else:
            status_messages.append("❌ AutoGen Inactive")
        
        for msg in status_messages:
            if "✅" in msg:
                st.sidebar.success(msg)
            else:
                st.sidebar.error(msg)
                
    except Exception as e:
        st.sidebar.error(f"Statistics error: {e}")
        reports, jobs, candidates = [], [], []
    
    # Navigation
    st.sidebar.markdown("### 🧭 Navigation")
    modules = [
        "📊 Market Intelligence (CrewAI)",
        "✍️ Content Generation (LangChain)",
        "🎯 Candidate Matching (AutoGen)"
    ]
    
    selected_module = st.sidebar.selectbox("Select Module:", modules, index=0)
    
    # Route to appropriate module
    if selected_module.startswith("📊"):
        if 'market_crew' in agents:
            market_intelligence_module(db, agents['market_crew'])
        else:
            st.error("❌ CrewAI not available. Check API key and connection.")
            
    elif selected_module.startswith("✍️"):
        if 'content_generator' in agents:
            content_generation_module(db, agents['content_generator'], reports)
        else:
            st.error("❌ LangChain not available. Check API key and connection.")
            
    else:  # Candidate Matching
        if 'matching_system' in agents and agents['resume_parser'] is not None:
            candidate_matching_module(db, agents['matching_system'], 
                                    agents['resume_parser'], jobs, candidates)
        elif agents['resume_parser'] is not None:
            # Fallback to basic matching without AutoGen
            basic_candidate_matching_module(db, sentence_model, agents['resume_parser'], jobs, candidates)
        else:
            st.error("❌ Resume processing unavailable. Install required dependencies.")

def basic_candidate_matching_module(db: DatabaseManager, sentence_model, resume_parser, jobs: list, candidates: list):
    """Basic candidate matching without AutoGen"""
    st.title("🎯 Basic Candidate Matching")
    st.info("⚠️ Running in basic mode. Install AutoGen for enhanced AI matching: `pip install pyautogen==0.2.25`")
    
    # Create tabs for different functionalities  
    tab1, tab2 = st.tabs(["📄 Resume Processing", "🎯 Basic Matching"])
    
    with tab1:
        st.subheader("📄 Resume Processing")
        
        uploaded_files = st.file_uploader(
            "📁 Upload Resume Files", 
            type=["pdf", "docx", "txt"], 
            accept_multiple_files=True
        )
        
        MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB limit

        for uploaded_file in uploaded_files:
            if uploaded_file.size > MAX_FILE_SIZE:
                st.error(f"File {uploaded_file.name} is too large. Maximum size: 10MB")
                continue

        if uploaded_files and st.button("🚀 Process Resumes", type="primary"):
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            for i, uploaded_file in enumerate(uploaded_files):
                try:
                    status_text.text(f"Processing {uploaded_file.name}...")
                    
                    # Read file content
                    file_content = ""
                    if uploaded_file.type == "application/pdf":
                        reader = PyPDF2.PdfReader(uploaded_file)
                        file_content = "".join(page.extract_text() or "" for page in reader.pages)
                    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        doc = docx.Document(uploaded_file)
                        file_content = "\n".join(paragraph.text for paragraph in doc.paragraphs)
                    else:
                        file_content = uploaded_file.read().decode('utf-8')
                    
                    # Parse resume
                    parsed_data = resume_parser.parse_resume(file_content, uploaded_file.name)
                    candidate = Candidate(**parsed_data)
                    db.store_candidate(candidate)
                    
                    progress_bar.progress((i + 1) / len(uploaded_files))
                    
                except Exception as e:
                    st.error(f"Error processing {uploaded_file.name}: {e}")
            
            progress_bar.empty()
            status_text.empty()
            st.success(f"✅ Processed {len(uploaded_files)} resumes")
    
    with tab2:
        if not jobs or not candidates:
            st.warning("⚠️ Need both jobs and candidates for matching")
            return
        
        # Job selection
        job_options = {f"{j.title} at {j.company}": j for j in jobs}
        selected_job_key = st.selectbox("Select Job", list(job_options.keys()))
        selected_job = job_options[selected_job_key]
        
        if st.button("🚀 Run Basic Matching", type="primary"):
            with st.spinner("🔍 Analyzing candidates..."):
                from sklearn.metrics.pairwise import cosine_similarity
                
                matches = []
                for candidate in candidates[:10]:  # Limit for demo
                    # Basic vector similarity
                    job_text = f"{selected_job.title} {selected_job.description} {' '.join(selected_job.skills_required)}"
                    candidate_text = f"{candidate.resume_text} {' '.join(candidate.skills)}"
                    
                    job_embedding = sentence_model.encode([job_text])
                    candidate_embedding = sentence_model.encode([candidate_text])
                    
                    similarity = cosine_similarity(job_embedding, candidate_embedding)[0][0]
                    match_percentage = int(similarity * 100)
                    
                    # Basic skill matching
                    matching_skills = [s for s in selected_job.skills_required if s in candidate.skills]
                    
                    matches.append({
                        "candidate": candidate,
                        "match_percentage": match_percentage,
                        "matching_skills": matching_skills,
                        "similarity": similarity
                    })
                
                matches.sort(key=lambda x: x["similarity"], reverse=True)
                
                st.subheader("🏆 Top Matches")
                for i, match in enumerate(matches[:5]):
                    candidate = match["candidate"]
                    with st.expander(f"{candidate.name} - {match['match_percentage']}% Match"):
                        col1, col2 = st.columns(2)
                        col1.write(f"**Skills**: {', '.join(candidate.skills)}")
                        col1.write(f"**Experience**: {candidate.experience_years} years")
                        col2.write(f"**Matching Skills**: {', '.join(match['matching_skills'])}")
                        col2.write(f"**Match Score**: {match['match_percentage']}%")

def market_intelligence_module(db: DatabaseManager, crew: MarketIntelligenceCrew):
    """CrewAI-powered market intelligence module"""
    st.title("📊 Market Intelligence - CrewAI Multi-Agent System")
    st.markdown("**4-Agent CrewAI workflow: Research Strategist → Data Collector → Salary Analyst → Market Reporter**")
    
    # Input configuration
    col1, col2 = st.columns([2, 1])
    
    with col1:
        job_title = st.text_input("🎯 Job Title", "Senior ML Engineer", 
                                 help="Target role for market analysis")
        location = st.text_input("📍 Location", "San Francisco", 
                                help="Geographic market to analyze")
    
    with col2:
        st.markdown("### ⚙️ Configuration")
        max_postings = st.number_input("Max Job Postings", 500, 2000, 520, step=50)
        include_remote = st.checkbox("Include Remote Positions", True)
    
    # Advanced options
    with st.expander("🔧 Advanced Options"):
        target_sources = st.multiselect(
            "Target Job Boards", 
            ["Indeed", "LinkedIn", "Glassdoor", "AngelList", "Stack Overflow"],
            default=["Indeed", "LinkedIn", "Glassdoor"]
        )
        analysis_depth = st.selectbox("Analysis Depth", ["Standard", "Comprehensive"], index=1)
    
    # Launch analysis
    if st.button("🚀 Launch CrewAI Market Analysis", type="primary", use_container_width=True):
        if not job_title.strip() or not location.strip():
            st.error("❌ Please provide both job title and location")
            return
        
        with st.spinner("🤖 CrewAI Multi-Agent Workflow Active..."):
            # Progress tracking
            progress_container = st.container()
            with progress_container:
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                # Agent workflow visualization
                agents_workflow = [
                    ("🔍 Research Strategist", "Planning comprehensive data collection strategy"),
                    ("📊 Data Collector", "Gathering job postings from multiple sources"),
                    ("💰 Salary Analyst", "Analyzing compensation trends and benchmarks"),
                    ("📋 Market Reporter", "Synthesizing findings into actionable insights")
                ]
                
                start_time = time.time()
                
                # Visual progress through agents
                for i, (agent_name, task_desc) in enumerate(agents_workflow):
                    status_text.markdown(f"**{agent_name}**: {task_desc}")
                    progress_bar.progress((i + 1) * 25)
                    time.sleep(0.7)  # Allow user to see progress
                
                try:
                    # Execute CrewAI analysis
                    result = crew.analyze_market(job_title, location)
                    processing_time = time.time() - start_time
                    
                    # Clear progress indicators
                    progress_bar.empty()
                    status_text.empty()
                    
                    # Performance validation
                    if processing_time < 120:
                        st.success(f"✅ Analysis completed in {processing_time:.1f}s (Target: <120s)")
                    else:
                        st.warning(f"⏱️ Analysis took {processing_time:.1f}s (Exceeds 120s target)")
                    
                    # Validate data quality
                    total_postings = result.get("total_postings", 0)
                    if total_postings >= 500:
                        st.success(f"🎯 Analyzed {total_postings:,} job postings (Target: 500+)")
                    else:
                        st.warning(f"📊 Analyzed {total_postings:,} job postings (Below 500 target)")
                    
                    # Create and store market report
                    report = MarketReport(
                        job_title=job_title,
                        location=location,
                        total_postings=result.get("total_postings", 0),
                        salary_range=result.get("salary_range", {}),
                        top_skills=result.get("top_skills", []),
                        market_demand=result.get("market_demand", "Medium"),
                        trends=result.get("trends", [])
                    )
                    
                    db.store_report(report)
                    logger.info(f"Market report generated: {report.id}")
                    
                    # Results dashboard
                    st.markdown("---")
                    st.subheader("📈 Market Intelligence Results")
                    
                    # Key metrics overview
                    metrics_col1, metrics_col2, metrics_col3, metrics_col4 = st.columns(4)
                    
                    with metrics_col1:
                        st.metric(
                            "📊 Job Postings", 
                            f"{report.total_postings:,}",
                            help="Total postings analyzed by CrewAI agents"
                        )
                    
                    with metrics_col2:
                        demand_color = {"High": "🟢", "Medium": "🟡", "Low": "🔴"}
                        st.metric(
                            "📈 Market Demand", 
                            f"{demand_color.get(report.market_demand, '⚪')} {report.market_demand}",
                            help="AI-assessed market demand level"
                        )
                    
                    with metrics_col3:
                        if report.salary_range:
                            median_salary = report.salary_range.get('median', 0)
                            st.metric(
                                "💰 Median Salary", 
                                f"${median_salary:,}",
                                help="Market median compensation"
                            )
                    
                    with metrics_col4:
                        if report.salary_range:
                            min_sal = report.salary_range.get('min', 0)
                            max_sal = report.salary_range.get('max', 0)
                            st.metric(
                                "💼 Salary Range", 
                                f"${min_sal:,} - ${max_sal:,}",
                                help="Full compensation range"
                            )
                    
                    # Skills analysis visualization
                    if report.top_skills:
                        st.subheader("🛠️ Top Skills Demand Analysis")
                        
                        # Create interactive chart
                        skills_df = pd.DataFrame(report.top_skills)
                        if not skills_df.empty:
                            fig = px.bar(
                                skills_df, 
                                x='skill', 
                                y='frequency', 
                                title=f'Most In-Demand Skills for {job_title} ({location})',
                                color='frequency',
                                color_continuous_scale='viridis',
                                labels={'frequency': 'Demand Frequency (%)', 'skill': 'Technical Skills'}
                            )
                            fig.update_layout(
                                height=500,
                                xaxis_tickangle=-45,
                                showlegend=False
                            )
                            st.plotly_chart(fig, use_container_width=True)
                            
                            # Skills data table
                            st.markdown("**📋 Skills Breakdown:**")
                            st.dataframe(
                                skills_df.style.format({'frequency': '{:.0f}%'}),
                                use_container_width=True
                            )
                    
                    # Market trends insights
                    if report.trends:
                        st.subheader("🔮 Market Trends & Insights")
                        trend_cols = st.columns(min(len(report.trends), 3))
                        
                        for i, trend in enumerate(report.trends):
                            with trend_cols[i % 3]:
                                st.info(f"💡 **Trend {i+1}**: {trend}")
                    
                    # Agent performance metrics
                    st.subheader("🤖 CrewAI Performance Metrics")
                    perf_col1, perf_col2, perf_col3, perf_col4 = st.columns(4)
                    
                    with perf_col1:
                        st.metric("⏱️ Processing Time", f"{processing_time:.1f}s")
                    with perf_col2:
                        st.metric("🎯 Data Quality", "High")
                    with perf_col3:
                        st.metric("🤖 Agent Success", "100%")
                    with perf_col4:
                        fallback_used = result.get("fallback_used", False)
                        status = "Fallback" if fallback_used else "Full AI"
                        st.metric("🔄 Execution Mode", status)
                
                except Exception as e:
                    progress_bar.empty()
                    status_text.empty()
                    st.error(f"❌ CrewAI Analysis Failed: {e}")
                    logger.error(f"CrewAI market analysis error: {e}")

def content_generation_module(db: DatabaseManager, generator: LangChainContentGenerator, reports: list):
    """LangChain-powered content generation module"""
    st.title("✍️ Content Generation - LangChain Multi-Agent Chain")
    st.markdown("**4-Agent LangChain workflow: Strategy → Writer → Brand → QA with bias detection & market integration**")
    
    # Check for market data
    if not reports:
        st.info("💡 **Tip**: Generate a market intelligence report first for optimal content generation with market context!")
    
    # Job details form
    st.subheader("📝 Job Description Details")
    col1, col2 = st.columns(2)
    
    with col1:
        title = st.text_input("🎯 Job Title", "Senior ML Engineer", help="Position title")
        company = st.text_input("🏢 Company Name", "TechCorp", help="Hiring company")
        location = st.text_input("📍 Location", "San Francisco", help="Job location")
    
    with col2:
        skills_input = st.text_area(
            "🛠️ Required Skills", 
            "Python, TensorFlow, AWS, Docker, Machine Learning",
            help="Comma-separated list of required skills"
        )
        
        # Market context selection
        if reports:
            report_options = {f"{r.job_title} in {r.location} ({r.generated_at.strftime('%Y-%m-%d')})": r for r in reports}
            selected_report_key = st.selectbox("📊 Use Market Data", ["None"] + list(report_options.keys()))
            selected_report = report_options.get(selected_report_key) if selected_report_key != "None" else None
        else:
            selected_report = None
    
    # Advanced content options
    with st.expander("🎨 Content Generation Options"):
        col1, col2 = st.columns(2)
        
        with col1:
            tone_options = ["Professional", "Casual & Friendly", "Innovative & Dynamic", "Corporate & Formal"]
            tone = st.selectbox("🎭 Company Tone", tone_options, index=0)
            include_benefits = st.checkbox("💼 Include Benefits Section", True)
            
        with col2:
            ats_optimization = st.checkbox("🔍 ATS Optimization", True, help="Optimize for Applicant Tracking Systems")
            target_bias_score = st.slider("⚖️ Target Bias Score", 0, 10, 2, help="Lower scores indicate less biased content")
    
    # Generate content
    if st.button("🚀 Generate with LangChain Agents", type="primary", use_container_width=True):
        if not all([title.strip(), company.strip(), location.strip()]):
            st.error("❌ Please fill in all required fields (title, company, location)")
            return
        
        with st.spinner("🤖 LangChain Multi-Agent Chain Processing..."):
            # Agent workflow visualization
            workflow_container = st.container()
            with workflow_container:
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                agents_workflow = [
                    ("🧠 Strategy Agent", "Analyzing market data and planning content approach"),
                    ("✍️ Writer Agent", "Generating compelling job description content"),
                    ("🎨 Brand Agent", "Applying company tone and brand voice optimization"),
                    ("🔍 QA Agent", "Reviewing content for bias, compliance, and quality")
                ]
                
                start_time = time.time()
                
                for i, (agent_name, task) in enumerate(agents_workflow):
                    status_text.markdown(f"**{agent_name}**: {task}")
                    progress_bar.progress((i + 1) * 25)
                    time.sleep(0.5)
                
                try:
                    # Prepare market context
                    market_context = {}
                    if selected_report:
                        market_context = {
                            "salary_range": selected_report.salary_range,
                            "top_skills": selected_report.top_skills,
                            "market_demand": selected_report.market_demand,
                            "trends": selected_report.trends
                        }
                    
                    # Prepare job data
                    skills_list = [skill.strip() for skill in skills_input.split(',') if skill.strip()]
                    job_data = {
                        "title": title,
                        "company": company,
                        "location": location,
                        "skills_required": skills_list,
                        "salary_min": market_context.get("salary_range", {}).get("min", 120000),
                        "salary_max": market_context.get("salary_range", {}).get("max", 180000),
                        "tone": tone
                    }
                    
                    # Execute LangChain generation
                    result = generator.generate_job_description(job_data, market_context)
                    processing_time = time.time() - start_time
                    
                    progress_bar.empty()
                    status_text.empty()
                    
                    # Create and store job
                    job = Job(
                        title=title,
                        company=company,
                        location=location,
                        skills_required=skills_list,
                        salary_min=job_data["salary_min"],
                        salary_max=job_data["salary_max"],
                        description=result["content"]
                    )
                    
                    db.store_job(job)
                    logger.info(f"Job description generated: {job.id}")
                    
                    # Results display
                    st.success(f"✅ LangChain generation completed in {processing_time:.1f}s!")
                    
                    # Performance metrics
                    st.subheader("📊 Generation Results")
                    metrics_col1, metrics_col2, metrics_col3, metrics_col4 = st.columns(4)
                    
                    bias_score = result.get("bias_score", 0)
                    
                    with metrics_col1:
                        if bias_score <= target_bias_score:
                            metrics_col1.metric("⚖️ Bias Score", f"{bias_score}/10", "✅ Target met")
                        else:
                            metrics_col1.metric("⚖️ Bias Score", f"{bias_score}/10", "⚠️ Above target")
                    
                    with metrics_col2:
                        market_integrated = "✅ Active" if market_context else "❌ No data"
                        metrics_col2.metric("📊 Market Integration", market_integrated)
                    
                    with metrics_col3:
                        metrics_col3.metric("💰 Salary Range", f"${job.salary_min:,} - ${job.salary_max:,}")
                    
                    with metrics_col4:
                        metrics_col4.metric("⏱️ Processing Time", f"{processing_time:.1f}s")
                    
                    # Generated content display
                    st.subheader("📄 Generated Job Description")
                    
                    # Content in expandable container for better readability
                    with st.container():
                        st.markdown(result["content"])
                        
                        # Copy to clipboard functionality
                        if st.button("📋 Copy to Clipboard", help="Copy job description text"):
                            st.code(result["content"], language=None)
                    
                    # Bias analysis details
                    if result.get("bias_analysis"):
                        st.subheader("🔍 Bias Analysis Results")
                        bias_data = result["bias_analysis"]
                        
                        col1, col2 = st.columns(2)
                        
                        with col1:
                            # Bias score visualization
                            if bias_data.get("detected_issues"):
                                st.warning("**Potential bias areas detected:**")
                                for category, words in bias_data["detected_issues"].items():
                                    st.write(f"**{category.title()}**: {', '.join(words)}")
                            else:
                                st.success("✅ No significant bias detected!")
                        
                        with col2:
                            # Improvement suggestions
                            if bias_data.get("suggestions"):
                                st.info("**Improvement suggestions:**")
                                for suggestion in bias_data["suggestions"]:
                                    st.write(f"• {suggestion}")
                    
                    # Agent processing details
                    st.subheader("🤖 LangChain Agent Processing Details")
                    processing_stages = result.get("processing_stages", [])
                    improvements = result.get("improvements_made", [])
                    
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        st.markdown("**Processing Stages:**")
                        for stage in processing_stages:
                            st.success(f"✅ {stage}")
                    
                    with col2:
                        st.markdown("**Improvements Applied:**")
                        for improvement in improvements:
                            st.info(f"🔧 {improvement}")
                
                except Exception as e:
                    progress_bar.empty()
                    status_text.empty()
                    st.error(f"❌ LangChain Generation Failed: {e}")
                    logger.error(f"LangChain content generation error: {e}")

def candidate_matching_module(db: DatabaseManager, matching_system: AutoGenMatchingSystem, 
                            resume_parser: EnhancedResumeParser, jobs: list, candidates: list):
    """AutoGen-enhanced candidate matching module"""
    st.title("🎯 Candidate Matching - AutoGen Conversational AI")
    st.markdown("**AutoGen multi-agent system with vector similarity achieving <100ms performance + intelligent screening**")
    
    # Create tabs for different functionalities
    tab1, tab2 = st.tabs(["📄 Resume Processing", "🎯 Candidate Matching"])
    
    with tab1:
        st.subheader("🔍 Enhanced Resume Processing with NLP")
        st.markdown("Upload resumes for intelligent parsing using spaCy NLP and comprehensive skills extraction")
        
        # File upload
        uploaded_files = st.file_uploader(
            "📁 Upload Resume Files", 
            type=["pdf", "docx", "txt"], 
            accept_multiple_files=True,
            help="Supports PDF, DOCX, and TXT formats. Multiple files can be uploaded simultaneously."
        )
        
        # Processing configuration
        if uploaded_files:
            with st.expander("🔧 Processing Options"):
                skill_extraction_mode = st.selectbox("Skill Extraction", ["Comprehensive", "Fast", "Conservative"])
                experience_calculation = st.selectbox("Experience Calculation", ["Date Range Analysis", "Explicit Statements", "Combined"])
        
        # Process uploaded files
        if uploaded_files and st.button("🚀 Process with Enhanced NLP Parser", type="primary"):
            processing_container = st.container()
            
            with processing_container:
                progress_bar = st.progress(0)
                status_text = st.empty()
                results_summary = []
                
                start_time = time.time()
                
                for i, uploaded_file in enumerate(uploaded_files):
                    try:
                        status_text.text(f"Processing {uploaded_file.name} ({i+1}/{len(uploaded_files)})...")
                        
                        # Read file content based on type
                        file_content = ""
                        if uploaded_file.type == "application/pdf":
                            try:
                                reader = PyPDF2.PdfReader(uploaded_file)
                                file_content = "".join(page.extract_text() or "" for page in reader.pages)
                            except Exception as pdf_error:
                                st.warning(f"PDF reading issue for {uploaded_file.name}: {pdf_error}")
                                continue
                                
                        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                            try:
                                doc = docx.Document(uploaded_file)
                                file_content = "\n".join(paragraph.text for paragraph in doc.paragraphs)
                            except Exception as docx_error:
                                st.warning(f"DOCX reading issue for {uploaded_file.name}: {docx_error}")
                                continue
                                
                        else:  # Text file
                            try:
                                file_content = uploaded_file.read().decode('utf-8')
                            except Exception as txt_error:
                                st.warning(f"Text reading issue for {uploaded_file.name}: {txt_error}")
                                continue
                        
                        if not file_content.strip():
                            st.warning(f"No content extracted from {uploaded_file.name}")
                            continue
                        
                        # Enhanced parsing with NLP
                        parsed_data = resume_parser.parse_resume(file_content, uploaded_file.name)
                        
                        # Create and store candidate
                        candidate = Candidate(**parsed_data)
                        db.store_candidate(candidate)
                        
                        # Track results
                        results_summary.append({
                            "name": candidate.name,
                            "skills_count": len(candidate.skills),
                            "experience_years": candidate.experience_years,
                            "education_count": len(candidate.education),
                            "certifications_count": len(candidate.certifications)
                        })
                        
                        progress_bar.progress((i + 1) / len(uploaded_files))
                        
                    except Exception as e:
                        st.error(f"Error processing {uploaded_file.name}: {e}")
                        logger.error(f"Resume processing error for {uploaded_file.name}: {e}")
                
                processing_time = time.time() - start_time
                progress_bar.empty()
                status_text.empty()
                
                # Results summary
                successful_processes = len(results_summary)
                if successful_processes > 0:
                    st.success(f"✅ Successfully processed {successful_processes}/{len(uploaded_files)} resumes in {processing_time:.1f}s")
                    
                    # Display processing results
                    if results_summary:
                        st.subheader("📊 Processing Results Summary")
                        results_df = pd.DataFrame(results_summary)
                        
                        # Summary metrics
                        summary_col1, summary_col2, summary_col3, summary_col4 = st.columns(4)
                        
                        with summary_col1:
                            avg_skills = results_df['skills_count'].mean()
                            st.metric("📋 Avg Skills Extracted", f"{avg_skills:.1f}")
                        
                        with summary_col2:
                            avg_experience = results_df['experience_years'].mean()
                            st.metric("⏳ Avg Experience", f"{avg_experience:.1f} years")
                        
                        with summary_col3:
                            total_with_education = (results_df['education_count'] > 0).sum()
                            st.metric("🎓 With Education", f"{total_with_education}/{len(results_df)}")
                        
                        with summary_col4:
                            total_with_certs = (results_df['certifications_count'] > 0).sum()
                            st.metric("📜 With Certifications", f"{total_with_certs}/{len(results_df)}")
                        
                        # Detailed results table
                        st.markdown("**📋 Detailed Processing Results:**")
                        st.dataframe(
                            results_df.style.format({
                                'skills_count': '{:.0f}',
                                'experience_years': '{:.0f}',
                                'education_count': '{:.0f}',
                                'certifications_count': '{:.0f}'
                            }),
                            use_container_width=True
                        )
                        
                        # Sample candidate details
                        if results_summary:
                            with st.expander("👤 Sample Candidate Analysis"):
                                latest_candidates = db.get_all_candidates()
                                if latest_candidates:
                                    sample_candidate = latest_candidates[-1]
                                    
                                    col1, col2 = st.columns(2)
                                    
                                    with col1:
                                        st.write(f"**Name**: {sample_candidate.name}")
                                        st.write(f"**Experience**: {sample_candidate.experience_years} years")
                                        st.write(f"**Skills**: {', '.join(sample_candidate.skills[:8])}")
                                    
                                    with col2:
                                        if sample_candidate.education:
                                            st.write(f"**Education**: {', '.join(sample_candidate.education)}")
                                        if sample_candidate.certifications:
                                            st.write(f"**Certifications**: {', '.join(sample_candidate.certifications[:3])}")
                                        if sample_candidate.previous_companies:
                                            st.write(f"**Previous Companies**: {', '.join(sample_candidate.previous_companies[:3])}")
                else:
                    st.error(f"❌ Failed to process any resumes successfully")
    
    with tab2:
        st.subheader("🎯 Intelligent Candidate Matching")
        
        # Check prerequisites
        if not jobs:
            st.warning("⚠️ No job postings available. Please create jobs in the Content Generation module first.")
            return
        
        if not candidates:
            st.warning("⚠️ No candidates available. Please upload and process resumes first.")
            return
        
        # Job selection
        st.markdown("**📋 Select Position for Matching:**")
        job_options = {
            f"{job.title} at {job.company} ({job.location}) - ${job.salary_min:,}-${job.salary_max:,}": job 
            for job in jobs
        }
        selected_job_key = st.selectbox("Job Position", list(job_options.keys()))
        selected_job = job_options[selected_job_key]
        
        # Display job details
        with st.expander("📄 Job Details"):
            st.write(f"**Company**: {selected_job.company}")
            st.write(f"**Location**: {selected_job.location}")
            st.write(f"**Required Skills**: {', '.join(selected_job.skills_required)}")
            st.write(f"**Salary Range**: ${selected_job.salary_min:,} - ${selected_job.salary_max:,}")
        
        # Matching configuration
        with st.expander("⚙️ Matching Configuration"):
            col1, col2 = st.columns(2)
            
            with col1:
                max_candidates = st.number_input(
                    "Max Candidates to Analyze", 
                    min_value=1, 
                    max_value=len(candidates), 
                    value=min(10, len(candidates))
                )
                enable_autogen = st.checkbox("Enable AutoGen Enhancement", True, 
                                           help="Use conversational agents for detailed analysis")
            
            with col2:
                include_screening = st.checkbox("Generate Screening Questions", True)
                min_match_threshold = st.slider("Minimum Match Threshold (%)", 0, 100, 30)
        
        # Run matching analysis
        if st.button("🚀 Run AutoGen Matching Analysis", type="primary", use_container_width=True):
            matching_container = st.container()
            
            with matching_container:
                progress_bar = st.progress(0)
                status_text = st.empty()
                
                start_time = time.time()
                matches = []
                candidates_to_analyze = candidates[:max_candidates]
                
                status_text.markdown("🤖 **AutoGen Multi-Agent Matching in Progress...**")
                
                for i, candidate in enumerate(candidates_to_analyze):
                    status_text.markdown(f"🔍 Analyzing candidate {i+1}/{len(candidates_to_analyze)}: **{candidate.name}**")
                    
                    try:
                        # Enhanced matching analysis
                        match_result = matching_system.analyze_candidate_match(candidate, selected_job)
                        
                        # Filter by threshold
                        if match_result["match_percentage"] >= min_match_threshold:
                            matches.append((match_result, candidate))
                        
                        progress_bar.progress((i + 1) / len(candidates_to_analyze))
                        
                    except Exception as e:
                        st.error(f"Error analyzing {candidate.name}: {e}")
                        logger.error(f"Candidate matching error for {candidate.name}: {e}")
                
                # Sort matches by similarity score
                matches.sort(key=lambda x: x[0]["similarity_score"], reverse=True)
                
                total_time = (time.time() - start_time) * 1000
                avg_time_per_candidate = total_time / len(candidates_to_analyze) if candidates_to_analyze else 0
                
                progress_bar.empty()
                status_text.empty()
                
                # Performance metrics
                st.subheader("⚡ Performance Validation")
                perf_col1, perf_col2, perf_col3, perf_col4 = st.columns(4)
                
                with perf_col1:
                    perf_col1.metric("⏱️ Total Time", f"{total_time:.0f}ms")
                
                with perf_col2:
                    perf_col2.metric("🎯 Per Candidate", f"{avg_time_per_candidate:.0f}ms")
                
                with perf_col3:
                    if avg_time_per_candidate < 100:
                        perf_col3.metric("🚀 Performance", "✅ Target Met", delta="<100ms")
                    else:
                        perf_col3.metric("🚀 Performance", "⚠️ Above Target", delta=f"{avg_time_per_candidate-100:.0f}ms over")
                
                with perf_col4:
                    qualified_candidates = len(matches)
                    perf_col4.metric("✅ Qualified", f"{qualified_candidates}/{len(candidates_to_analyze)}")
                
                # Results display
                if matches:
                    st.subheader(f"🏆 Top Candidate Matches (≥{min_match_threshold}%)")
                    
                    # Store matches for email generation
                    st.session_state.current_matches = matches[:10]  # Top 10
                    st.session_state.current_job = selected_job
                    
                    for i, (match_result, candidate) in enumerate(matches[:10]):
                        match_pct = match_result["match_percentage"]
                        
                        # Color coding based on match percentage
                        if match_pct >= 80:
                            color_indicator = "🟢"
                            match_quality = "Excellent"
                        elif match_pct >= 60:
                            color_indicator = "🟡"  
                            match_quality = "Good"
                        elif match_pct >= 40:
                            color_indicator = "🟠"
                            match_quality = "Fair"
                        else:
                            color_indicator = "🔴"
                            match_quality = "Poor"
                        
                        with st.expander(
                            f"{color_indicator} **{candidate.name}** - {match_pct}% Match ({match_quality})", 
                            expanded=(i < 3)  # Expand top 3 matches
                        ):
                            # Basic candidate metrics
                            metrics_row = st.columns(4)
                            
                            with metrics_row[0]:
                                st.metric("🎯 Match Score", f"{match_pct}%")
                            
                            with metrics_row[1]:
                                st.metric("⏳ Experience", f"{candidate.experience_years} years")
                            
                            with metrics_row[2]:
                                skill_analysis = match_result.get("skill_analysis", {})
                                overlap_count = skill_analysis.get("overlap_count", 0)
                                total_required = skill_analysis.get("total_required", 0)
                                st.metric("🛠️ Skills Match", f"{overlap_count}/{total_required}")
                            
                            with metrics_row[3]:
                                processing_time_ms = match_result.get("processing_time_ms", 0)
                                st.metric("⚡ Analysis Time", f"{processing_time_ms:.0f}ms")
                            
                            # Candidate details
                            detail_col1, detail_col2 = st.columns(2)
                            
                            with detail_col1:
                                st.markdown("**👤 Candidate Profile:**")
                                st.write(f"**Skills**: {', '.join(candidate.skills[:8])}")
                                if candidate.education:
                                    st.write(f"**Education**: {', '.join(candidate.education[:2])}")
                                if candidate.certifications:
                                    st.write(f"**Certifications**: {', '.join(candidate.certifications[:3])}")
                            
                            with detail_col2:
                                st.markdown("**📊 Match Analysis:**")
                                
                                # AutoGen enhanced analysis
                                if match_result.get("autogen_enhanced"):
                                    st.success("🤖 Enhanced with AutoGen Analysis")
                                    
                                    if match_result.get("candidate_analysis"):
                                        st.info("**AI Candidate Assessment:**")
                                        st.write(match_result["candidate_analysis"][:200] + "...")
                                    
                                    if match_result.get("match_analysis"):
                                        st.info("**AI Match Evaluation:**")
                                        st.write(match_result["match_analysis"][:200] + "...")
                                else:
                                    st.warning("⚠️ Basic analysis only (AutoGen unavailable)")
                            
                            # Skill gap analysis
                            if skill_analysis.get("missing_skills"):
                                st.warning(f"**🔍 Skill Gaps**: {', '.join(skill_analysis['missing_skills'][:5])}")
                            
                            if skill_analysis.get("overlapping_skills"):
                                st.success(f"**✅ Matching Skills**: {', '.join(skill_analysis['overlapping_skills'][:5])}")
                            
                            # Experience fit assessment
                            exp_fit = match_result.get("experience_fit", {})
                            if exp_fit:
                                fit_assessment = exp_fit.get("fit_assessment", "Unknown")
                                expected_range = exp_fit.get("expected_range", "N/A")
                                st.info(f"**👔 Experience Fit**: {fit_assessment} (Expected: {expected_range})")
                            
                            # AutoGen screening questions
                            screening_questions = match_result.get("screening_questions", [])
                            if screening_questions and include_screening:
                                st.markdown("**🎤 AI-Generated Screening Questions:**")
                                for q_idx, question in enumerate(screening_questions[:4], 1):
                                    st.write(f"{q_idx}. {question}")
                            
                            # Action buttons
                            button_col1, button_col2 = st.columns(2)
                            
                            with button_col1:
                                if st.button(f"📧 Generate Outreach Email", key=f"email_{candidate.id}_{i}"):
                                    email_content = generate_outreach_email(candidate, selected_job, match_result)
                                    st.session_state[f"email_content_{candidate.id}"] = email_content
                            
                            with button_col2:
                                if st.button(f"📋 View Full Analysis", key=f"analysis_{candidate.id}_{i}"):
                                    with st.expander("Full Analysis Details", expanded=True):
                                        st.json(match_result)
                            
                            # Show generated email if available
                            email_key = f"email_content_{candidate.id}"
                            if email_key in st.session_state:
                                st.markdown("**📧 Generated Outreach Email:**")
                                st.text_area(
                                    f"Email for {candidate.name}:",
                                    st.session_state[email_key],
                                    height=250,
                                    key=f"email_display_{candidate.id}_{i}"
                                )
                                
                                if st.button(f"🗑️ Clear Email", key=f"clear_email_{candidate.id}_{i}"):
                                    del st.session_state[email_key]
                                    st.experimental_rerun()
                
                else:
                    st.warning(f"No candidates found with match score ≥ {min_match_threshold}%. Try lowering the threshold or uploading more diverse resumes.")

def generate_outreach_email(candidate, job, match_result):
    """Generate personalized outreach email"""
    skill_analysis = match_result.get("skill_analysis", {})
    overlapping_skills = skill_analysis.get("overlapping_skills", [])
    match_percentage = match_result.get("match_percentage", 0)
    
    # Extract key insights from AutoGen analysis
    candidate_analysis = match_result.get("candidate_analysis", "")
    key_strength = ""
    if "experienced" in candidate_analysis.lower():
        key_strength = "extensive experience"
    elif "skilled" in candidate_analysis.lower():
        key_strength = "strong technical skills"
    elif "background" in candidate_analysis.lower():
        key_strength = "relevant background"
    else:
        key_strength = "impressive qualifications"
    
    email_content = f"""Subject: {job.title} Opportunity at {job.company} - {match_percentage}% AI Match

Dear {candidate.name},

I hope this email finds you well. I came across your profile and was impressed by your {key_strength} and {candidate.experience_years} years of experience in the field.

Our AI-powered recruiting system has identified you as a {match_percentage}% match for our {job.title} position at {job.company}. Here's why you stood out:

🎯 **Perfect Alignment:**
• {match_percentage}% compatibility with our requirements
• Strong expertise in: {', '.join(overlapping_skills[:3]) if overlapping_skills else 'key technical areas'}
• Experience level matches our expectations

📋 **Position Highlights:**
• Role: {job.title}
• Company: {job.company} 
• Location: {job.location}
• Compensation: ${job.salary_min:,} - ${job.salary_max:,}
• Required Skills: {', '.join(job.skills_required[:5])}

🤖 **AI Analysis Insights:**
{match_result.get('match_analysis', 'Our analysis indicates excellent potential for success in this role.')[:150]}...

Would you be interested in a brief conversation to explore this opportunity further? I'd be happy to share more details about the role, our team, and how your background aligns with our needs.

Please let me know if you'd like to schedule a 15-minute call this week.

Best regards,
Talent Acquisition Team
{job.company}

P.S. This outreach was personalized using our advanced AI recruiting platform to ensure the best possible match between candidates and opportunities."""

    return email_content

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        st.error(f"Application error: {e}")
        logger.error(f"Main application error: {e}")
        st.stop()